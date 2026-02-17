"""Document text extraction service for PDF and DOCX files."""

from io import BytesIO
from typing import Literal

import pdfplumber
from docx import Document
from pdfminer.pdfparser import PDFSyntaxError


class DocumentExtractionError(Exception):
    """Raised when document extraction fails."""

    pass


class ExtractionResult:
    """Result of document text extraction."""

    def __init__(
        self,
        raw_content: str,
        source_filename: str,
        file_type: Literal["pdf", "docx"],
        page_count: int | None,
        word_count: int,
        warnings: list[str] | None = None,
    ):
        self.raw_content = raw_content
        self.source_filename = source_filename
        self.file_type = file_type
        self.page_count = page_count
        self.word_count = word_count
        self.warnings = warnings or []


def extract_text_from_pdf(file_bytes: bytes, filename: str) -> ExtractionResult:
    """Extract text from a PDF file.

    Args:
        file_bytes: Raw bytes of the PDF file
        filename: Original filename for metadata

    Returns:
        ExtractionResult with extracted text and metadata

    Raises:
        DocumentExtractionError: If extraction fails
    """
    warnings: list[str] = []

    try:
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                raise DocumentExtractionError("PDF contains no pages")

            page_count = len(pdf.pages)
            text_parts: list[str] = []

            for i, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    else:
                        warnings.append(f"Page {i + 1} contained no extractable text")
                except Exception as e:
                    warnings.append(f"Failed to extract text from page {i + 1}: {str(e)}")

            raw_content = "\n\n".join(text_parts)

            if not raw_content.strip():
                raise DocumentExtractionError(
                    "Could not extract any text from PDF. "
                    "The file may be image-based or password-protected."
                )

            word_count = len(raw_content.split())

            return ExtractionResult(
                raw_content=raw_content,
                source_filename=filename,
                file_type="pdf",
                page_count=page_count,
                word_count=word_count,
                warnings=warnings,
            )

    except PDFSyntaxError:
        raise DocumentExtractionError(
            "Invalid or corrupted PDF file. Please ensure the file is a valid PDF."
        )
    except DocumentExtractionError:
        raise
    except Exception as e:
        raise DocumentExtractionError(f"Failed to process PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes, filename: str) -> ExtractionResult:
    """Extract text from a DOCX file.

    Args:
        file_bytes: Raw bytes of the DOCX file
        filename: Original filename for metadata

    Returns:
        ExtractionResult with extracted text and metadata

    Raises:
        DocumentExtractionError: If extraction fails
    """
    warnings: list[str] = []

    try:
        doc = Document(BytesIO(file_bytes))

        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        raw_content = "\n\n".join(paragraphs)

        if not raw_content.strip():
            raise DocumentExtractionError(
                "Could not extract any text from DOCX. The document appears to be empty."
            )

        word_count = len(raw_content.split())

        return ExtractionResult(
            raw_content=raw_content,
            source_filename=filename,
            file_type="docx",
            page_count=None,
            word_count=word_count,
            warnings=warnings,
        )

    except DocumentExtractionError:
        raise
    except Exception as e:
        raise DocumentExtractionError(f"Failed to process DOCX: {str(e)}")


def extract_text(
    file_bytes: bytes, filename: str, content_type: str
) -> ExtractionResult:
    """Extract text from a document based on its content type.

    Args:
        file_bytes: Raw bytes of the file
        filename: Original filename
        content_type: MIME type of the file

    Returns:
        ExtractionResult with extracted text and metadata

    Raises:
        DocumentExtractionError: If extraction fails or file type is unsupported
    """
    if content_type == "application/pdf":
        return extract_text_from_pdf(file_bytes, filename)
    elif (
        content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return extract_text_from_docx(file_bytes, filename)
    else:
        raise DocumentExtractionError(
            f"Unsupported file type: {content_type}. Only PDF and DOCX files are supported."
        )
