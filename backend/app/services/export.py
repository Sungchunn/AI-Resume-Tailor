import io
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem


class ExportService:
    """Service for exporting tailored resumes to different formats."""

    def generate_plain_text(self, tailored_content: dict[str, Any]) -> str:
        """Generate plain text version of a tailored resume."""
        lines = []

        # Summary
        if summary := tailored_content.get("summary"):
            lines.append("SUMMARY")
            lines.append("-" * 50)
            lines.append(summary)
            lines.append("")

        # Experience
        if experience := tailored_content.get("experience", []):
            lines.append("EXPERIENCE")
            lines.append("-" * 50)
            for job in experience:
                lines.append(f"{job.get('title', '')} | {job.get('company', '')}")
                location = job.get("location", "")
                dates = f"{job.get('start_date', '')} - {job.get('end_date', '')}"
                lines.append(f"{location} | {dates}")
                for bullet in job.get("bullets", []):
                    lines.append(f"  - {bullet}")
                lines.append("")

        # Skills
        if skills := tailored_content.get("skills", []):
            lines.append("SKILLS")
            lines.append("-" * 50)
            lines.append(", ".join(skills))
            lines.append("")

        # Highlights
        if highlights := tailored_content.get("highlights", []):
            lines.append("KEY HIGHLIGHTS")
            lines.append("-" * 50)
            for highlight in highlights:
                lines.append(f"  - {highlight}")
            lines.append("")

        return "\n".join(lines)

    def generate_docx(self, tailored_content: dict[str, Any]) -> bytes:
        """Generate DOCX version of a tailored resume."""
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Summary
        if summary := tailored_content.get("summary"):
            heading = doc.add_heading("Summary", level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph(summary)
            doc.add_paragraph()

        # Experience
        if experience := tailored_content.get("experience", []):
            doc.add_heading("Experience", level=1)
            for job in experience:
                # Job title and company
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{job.get('title', '')}")
                title_run.bold = True
                title_para.add_run(f" | {job.get('company', '')}")

                # Location and dates
                location = job.get("location", "")
                dates = f"{job.get('start_date', '')} - {job.get('end_date', '')}"
                date_para = doc.add_paragraph(f"{location} | {dates}")
                date_para.runs[0].italic = True

                # Bullets
                for bullet in job.get("bullets", []):
                    bullet_para = doc.add_paragraph(bullet, style="List Bullet")

                doc.add_paragraph()

        # Skills
        if skills := tailored_content.get("skills", []):
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(skills))
            doc.add_paragraph()

        # Highlights
        if highlights := tailored_content.get("highlights", []):
            doc.add_heading("Key Highlights", level=1)
            for highlight in highlights:
                doc.add_paragraph(highlight, style="List Bullet")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_pdf(self, tailored_content: dict[str, Any]) -> bytes:
        """Generate PDF version of a tailored resume."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            "Heading",
            parent=styles["Heading1"],
            fontSize=14,
            spaceAfter=6,
            textColor=colors.HexColor("#1f2937"),
        ))
        styles.add(ParagraphStyle(
            "JobTitle",
            parent=styles["Normal"],
            fontSize=12,
            spaceAfter=3,
            textColor=colors.HexColor("#111827"),
            fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            "JobDetails",
            parent=styles["Normal"],
            fontSize=10,
            spaceAfter=6,
            textColor=colors.HexColor("#6b7280"),
            fontName="Helvetica-Oblique",
        ))
        styles.add(ParagraphStyle(
            "BodyText",
            parent=styles["Normal"],
            fontSize=11,
            spaceAfter=12,
            textColor=colors.HexColor("#374151"),
        ))
        styles.add(ParagraphStyle(
            "Bullet",
            parent=styles["Normal"],
            fontSize=10,
            leftIndent=20,
            spaceAfter=4,
            textColor=colors.HexColor("#374151"),
        ))

        elements = []

        # Summary
        if summary := tailored_content.get("summary"):
            elements.append(Paragraph("SUMMARY", styles["Heading"]))
            elements.append(Paragraph(summary, styles["BodyText"]))
            elements.append(Spacer(1, 12))

        # Experience
        if experience := tailored_content.get("experience", []):
            elements.append(Paragraph("EXPERIENCE", styles["Heading"]))
            for job in experience:
                title = job.get("title", "")
                company = job.get("company", "")
                elements.append(Paragraph(f"{title} | {company}", styles["JobTitle"]))

                location = job.get("location", "")
                dates = f"{job.get('start_date', '')} - {job.get('end_date', '')}"
                elements.append(Paragraph(f"{location} | {dates}", styles["JobDetails"]))

                for bullet in job.get("bullets", []):
                    # Escape special characters for ReportLab
                    safe_bullet = bullet.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    elements.append(Paragraph(f"• {safe_bullet}", styles["Bullet"]))

                elements.append(Spacer(1, 8))

        # Skills
        if skills := tailored_content.get("skills", []):
            elements.append(Paragraph("SKILLS", styles["Heading"]))
            skills_text = ", ".join(skills)
            elements.append(Paragraph(skills_text, styles["BodyText"]))
            elements.append(Spacer(1, 12))

        # Highlights
        if highlights := tailored_content.get("highlights", []):
            elements.append(Paragraph("KEY HIGHLIGHTS", styles["Heading"]))
            for highlight in highlights:
                safe_highlight = highlight.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                elements.append(Paragraph(f"• {safe_highlight}", styles["Bullet"]))

        doc.build(elements)
        buffer.seek(0)
        return buffer.getvalue()


def get_export_service() -> ExportService:
    """Get the export service instance."""
    return ExportService()
