"""
Export Service for Resume Generation.

Converts tailored resume content to PDF, DOCX, and plain text formats.
PDF generation uses WeasyPrint with Jinja2 templates.
"""

import io
from dataclasses import dataclass
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from app.services.export.html_to_document import ExportOptions, WEASYPRINT_AVAILABLE
from app.services.export.template_renderer import (
    ExportStyle,
    get_template_renderer,
)

# WeasyPrint import (conditional for environments without system deps)
if WEASYPRINT_AVAILABLE:
    from weasyprint import HTML
    from weasyprint.text.fonts import FontConfiguration


@dataclass
class PDFResult:
    """Result of PDF generation with metadata."""

    content: bytes
    page_count: int
    overflows: bool = False


class ExportService:
    """Service for exporting tailored resumes to different formats."""

    def __init__(self):
        self._template_renderer = get_template_renderer()
        self._font_config = FontConfiguration() if WEASYPRINT_AVAILABLE else None

    def generate_plain_text(self, tailored_content: dict[str, Any]) -> str:
        """Generate plain text version of a tailored resume."""
        lines = []

        # Summary
        if summary := tailored_content.get("summary"):
            lines.append("SUMMARY")
            lines.append("-" * 50)
            lines.append(summary)
            lines.append("")

        # Experience
        if experience := tailored_content.get("experience", []):
            lines.append("EXPERIENCE")
            lines.append("-" * 50)
            for job in experience:
                lines.append(f"{job.get('title', '')} | {job.get('company', '')}")
                location = job.get("location", "")
                dates = f"{job.get('start_date', '')} - {job.get('end_date', '')}"
                lines.append(f"{location} | {dates}")
                for bullet in job.get("bullets", []):
                    lines.append(f"  - {bullet}")
                lines.append("")

        # Skills
        if skills := tailored_content.get("skills", []):
            lines.append("SKILLS")
            lines.append("-" * 50)
            lines.append(", ".join(skills))
            lines.append("")

        # Highlights
        if highlights := tailored_content.get("highlights", []):
            lines.append("KEY HIGHLIGHTS")
            lines.append("-" * 50)
            for highlight in highlights:
                lines.append(f"  - {highlight}")
            lines.append("")

        return "\n".join(lines)

    def generate_docx(
        self,
        tailored_content: dict[str, Any],
        options: ExportOptions | None = None,
    ) -> bytes:
        """Generate DOCX version of a tailored resume."""
        if options is None:
            options = ExportOptions()

        doc = Document()

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(options.margin_top)
            section.bottom_margin = Inches(options.margin_bottom)
            section.left_margin = Inches(options.margin_left)
            section.right_margin = Inches(options.margin_right)

        # Set default font
        style = doc.styles["Normal"]
        style.font.name = options.font_family
        style.font.size = Pt(options.font_size)
        style.paragraph_format.line_spacing = options.line_spacing

        # Summary
        if summary := tailored_content.get("summary"):
            heading = doc.add_heading("Summary", level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph(summary)
            doc.add_paragraph()

        # Experience
        if experience := tailored_content.get("experience", []):
            doc.add_heading("Experience", level=1)
            for job in experience:
                # Job title and company
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{job.get('title', '')}")
                title_run.bold = True
                title_para.add_run(f" | {job.get('company', '')}")

                # Location and dates
                location = job.get("location", "")
                dates = f"{job.get('start_date', '')} - {job.get('end_date', '')}"
                date_para = doc.add_paragraph(f"{location} | {dates}")
                date_para.runs[0].italic = True

                # Bullets
                for bullet in job.get("bullets", []):
                    doc.add_paragraph(bullet, style="List Bullet")

                doc.add_paragraph()

        # Skills
        if skills := tailored_content.get("skills", []):
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(skills))
            doc.add_paragraph()

        # Highlights
        if highlights := tailored_content.get("highlights", []):
            doc.add_heading("Key Highlights", level=1)
            for highlight in highlights:
                doc.add_paragraph(highlight, style="List Bullet")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_pdf(
        self,
        tailored_content: dict[str, Any],
        options: ExportOptions | None = None,
        contact: dict[str, Any] | None = None,
    ) -> PDFResult:
        """
        Generate PDF version of a tailored resume using WeasyPrint.

        Args:
            tailored_content: Resume content dict with sections
            options: Export customization options
            contact: Optional contact information dict

        Returns:
            PDFResult with content bytes, page_count, and overflow flag
        """
        if not WEASYPRINT_AVAILABLE:
            raise RuntimeError(
                "PDF export requires WeasyPrint system dependencies. "
                "Install with: brew install pango (macOS) or "
                "apt-get install libpango-1.0-0 libpangocairo-1.0-0 (Ubuntu)"
            )

        if options is None:
            options = ExportOptions()

        # Convert ExportOptions to ExportStyle
        style = ExportStyle(
            template=options.template.value if hasattr(options.template, "value") else str(options.template),
            font_family=options.font_family,
            font_size=options.font_size,
            margin_top=options.margin_top,
            margin_bottom=options.margin_bottom,
            margin_left=options.margin_left,
            margin_right=options.margin_right,
            line_spacing=options.line_spacing,
            page_size=options.page_size,
        )

        # Normalize content to unified structure
        normalized = self._template_renderer.normalize_tailored_content(
            tailored_content, contact=contact
        )

        # Render HTML using template
        html_content = self._template_renderer.render(normalized, style)

        # Generate PDF using WeasyPrint
        html = HTML(string=html_content)
        document = html.render(font_config=self._font_config)

        # Get page count
        page_count = len(document.pages)

        # Check for overflow (content doesn't fit on target pages)
        # For now, we consider overflow if > 1 page and user expected 1
        overflows = page_count > 1

        # Write PDF to bytes
        buffer = io.BytesIO()
        document.write_pdf(buffer)
        buffer.seek(0)

        return PDFResult(
            content=buffer.getvalue(),
            page_count=page_count,
            overflows=overflows,
        )

    # Async wrapper methods for workshop router compatibility
    async def export_pdf(
        self,
        content: dict[str, Any],
        template: str = "classic",
        contact: dict[str, Any] | None = None,
    ) -> PDFResult:
        """Export resume content as PDF bytes with metadata."""
        from app.services.export.html_to_document import StyleTemplate

        try:
            style_template = StyleTemplate(template)
        except ValueError:
            style_template = StyleTemplate.CLASSIC

        options = ExportOptions(template=style_template)
        return self.generate_pdf(content, options=options, contact=contact)

    async def export_docx(
        self,
        content: dict[str, Any],
        template: str = "default",
    ) -> bytes:
        """Export resume content as DOCX bytes."""
        return self.generate_docx(content)

    async def export_txt(
        self,
        content: dict[str, Any],
        template: str = "default",
    ) -> str:
        """Export resume content as plain text."""
        return self.generate_plain_text(content)

    async def export_json(
        self,
        content: dict[str, Any],
        template: str = "default",
    ) -> str:
        """Export resume content as JSON string."""
        import json

        return json.dumps(content, indent=2)


def get_export_service() -> ExportService:
    """Get the export service instance."""
    return ExportService()
