"""
HTML to Document Export Service.

Converts TipTap HTML content to PDF and DOCX formats with style templates.

Note: PDF export requires WeasyPrint system dependencies:
  - macOS: brew install pango
  - Ubuntu: apt-get install libpango-1.0-0 libpangocairo-1.0-0
"""

import io
import re
from dataclasses import dataclass
from enum import Enum
from typing import Literal

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# WeasyPrint requires system dependencies (pango, gobject)
# Import conditionally to allow DOCX export even without PDF support
try:
    from weasyprint import HTML
    from weasyprint.text.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False
    HTML = None
    FontConfiguration = None


class StyleTemplate(str, Enum):
    """Available resume style templates."""
    CLASSIC = "classic"
    MODERN = "modern"
    MINIMAL = "minimal"


@dataclass
class ExportOptions:
    """Export customization options."""
    template: StyleTemplate = StyleTemplate.CLASSIC
    font_family: str = "Arial"
    font_size: int = 11  # Base font size in points
    margin_top: float = 0.75  # inches
    margin_bottom: float = 0.75
    margin_left: float = 0.75
    margin_right: float = 0.75
    line_spacing: float = 1.15
    page_size: str = "letter"  # "letter" or "a4"


# Page size CSS values
PAGE_SIZE_CSS = {
    "letter": "letter",
    "a4": "210mm 297mm",
}


# CSS Templates for PDF generation
TEMPLATE_CSS = {
    StyleTemplate.CLASSIC: """
        @page {{
            size: {page_size};
            margin: {margin_top}in {margin_right}in {margin_bottom}in {margin_left}in;
        }}
        body {{
            font-family: {font_family}, Times New Roman, serif;
            font-size: {font_size}pt;
            line-height: {line_spacing};
            color: #1a1a1a;
            margin: 0;
            padding: 0;
        }}
        h1 {{
            font-size: {h1_size}pt;
            font-weight: bold;
            color: #000;
            margin-top: 0;
            margin-bottom: 12pt;
            border-bottom: 2px solid #333;
            padding-bottom: 4pt;
        }}
        h2 {{
            font-size: {h2_size}pt;
            font-weight: bold;
            color: #333;
            margin-top: 16pt;
            margin-bottom: 8pt;
            text-transform: uppercase;
            letter-spacing: 1px;
            border-bottom: 1px solid #666;
            padding-bottom: 2pt;
        }}
        h3 {{
            font-size: {h3_size}pt;
            font-weight: bold;
            color: #444;
            margin-top: 10pt;
            margin-bottom: 4pt;
        }}
        p {{
            margin-top: 0;
            margin-bottom: 8pt;
        }}
        ul, ol {{
            margin-top: 4pt;
            margin-bottom: 8pt;
            padding-left: 20pt;
        }}
        li {{
            margin-bottom: 4pt;
        }}
        strong {{
            font-weight: bold;
        }}
        em {{
            font-style: italic;
        }}
        a {{
            color: #0066cc;
            text-decoration: none;
        }}
        blockquote {{
            margin: 8pt 0;
            padding-left: 12pt;
            border-left: 3px solid #ccc;
            color: #555;
            font-style: italic;
        }}
    """,
    StyleTemplate.MODERN: """
        @page {{
            size: {page_size};
            margin: {margin_top}in {margin_right}in {margin_bottom}in {margin_left}in;
        }}
        body {{
            font-family: {font_family}, Helvetica, sans-serif;
            font-size: {font_size}pt;
            line-height: {line_spacing};
            color: #2d2d2d;
            margin: 0;
            padding: 0;
        }}
        h1 {{
            font-size: {h1_size}pt;
            font-weight: 300;
            color: #1a365d;
            margin-top: 0;
            margin-bottom: 16pt;
            letter-spacing: 2px;
        }}
        h2 {{
            font-size: {h2_size}pt;
            font-weight: 600;
            color: #2563eb;
            margin-top: 20pt;
            margin-bottom: 10pt;
            text-transform: uppercase;
            letter-spacing: 1.5px;
        }}
        h3 {{
            font-size: {h3_size}pt;
            font-weight: 500;
            color: #1e40af;
            margin-top: 12pt;
            margin-bottom: 6pt;
        }}
        p {{
            margin-top: 0;
            margin-bottom: 10pt;
        }}
        ul, ol {{
            margin-top: 6pt;
            margin-bottom: 10pt;
            padding-left: 24pt;
        }}
        li {{
            margin-bottom: 6pt;
        }}
        li::marker {{
            color: #2563eb;
        }}
        strong {{
            font-weight: 600;
            color: #1e293b;
        }}
        em {{
            font-style: italic;
            color: #475569;
        }}
        a {{
            color: #2563eb;
            text-decoration: none;
        }}
        blockquote {{
            margin: 12pt 0;
            padding: 12pt 16pt;
            background: #f1f5f9;
            border-left: 4px solid #2563eb;
            border-radius: 0 4px 4px 0;
        }}
    """,
    StyleTemplate.MINIMAL: """
        @page {{
            size: {page_size};
            margin: {margin_top}in {margin_right}in {margin_bottom}in {margin_left}in;
        }}
        body {{
            font-family: {font_family}, system-ui, sans-serif;
            font-size: {font_size}pt;
            line-height: {line_spacing};
            color: #18181b;
            margin: 0;
            padding: 0;
        }}
        h1 {{
            font-size: {h1_size}pt;
            font-weight: 500;
            color: #09090b;
            margin-top: 0;
            margin-bottom: 8pt;
        }}
        h2 {{
            font-size: {h2_size}pt;
            font-weight: 500;
            color: #27272a;
            margin-top: 18pt;
            margin-bottom: 6pt;
        }}
        h3 {{
            font-size: {h3_size}pt;
            font-weight: 500;
            color: #3f3f46;
            margin-top: 10pt;
            margin-bottom: 4pt;
        }}
        p {{
            margin-top: 0;
            margin-bottom: 6pt;
        }}
        ul, ol {{
            margin-top: 4pt;
            margin-bottom: 6pt;
            padding-left: 18pt;
        }}
        li {{
            margin-bottom: 3pt;
        }}
        strong {{
            font-weight: 500;
        }}
        em {{
            font-style: italic;
        }}
        a {{
            color: #18181b;
            text-decoration: underline;
        }}
        blockquote {{
            margin: 8pt 0;
            padding-left: 12pt;
            border-left: 2px solid #d4d4d8;
            color: #52525b;
        }}
    """,
}


class HTMLToDocumentService:
    """Service for converting HTML content to PDF and DOCX documents."""

    def __init__(self):
        self._font_config = FontConfiguration() if WEASYPRINT_AVAILABLE else None

    def _get_css(self, options: ExportOptions) -> str:
        """Generate CSS for the given template and options."""
        template_css = TEMPLATE_CSS.get(options.template, TEMPLATE_CSS[StyleTemplate.CLASSIC])
        page_size_css = PAGE_SIZE_CSS.get(options.page_size, "letter")

        return template_css.format(
            font_family=options.font_family,
            font_size=options.font_size,
            h1_size=options.font_size + 8,
            h2_size=options.font_size + 4,
            h3_size=options.font_size + 2,
            margin_top=options.margin_top,
            margin_bottom=options.margin_bottom,
            margin_left=options.margin_left,
            margin_right=options.margin_right,
            line_spacing=options.line_spacing,
            page_size=page_size_css,
        )

    def _wrap_html(self, html_content: str, options: ExportOptions) -> str:
        """Wrap HTML content in a full document structure."""
        css = self._get_css(options)
        return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        {css}
    </style>
</head>
<body>
    {html_content}
</body>
</html>
"""

    def generate_pdf(
        self,
        html_content: str,
        options: ExportOptions | None = None,
    ) -> bytes:
        """
        Generate PDF from HTML content.

        Args:
            html_content: TipTap HTML content
            options: Export customization options

        Returns:
            PDF file as bytes

        Raises:
            RuntimeError: If WeasyPrint is not available
        """
        if not WEASYPRINT_AVAILABLE:
            raise RuntimeError(
                "PDF export requires WeasyPrint system dependencies. "
                "Install with: brew install pango (macOS) or "
                "apt-get install libpango-1.0-0 libpangocairo-1.0-0 (Ubuntu)"
            )

        if options is None:
            options = ExportOptions()

        full_html = self._wrap_html(html_content, options)

        # Generate PDF using WeasyPrint
        html = HTML(string=full_html)
        buffer = io.BytesIO()
        html.write_pdf(buffer, font_config=self._font_config)
        buffer.seek(0)
        return buffer.getvalue()

    def render_page_count(
        self,
        html_content: str,
        options: ExportOptions | None = None,
    ) -> int:
        """
        Render HTML and return page count without writing PDF.

        Args:
            html_content: TipTap HTML content
            options: Export customization options

        Returns:
            Number of pages the content would occupy

        Raises:
            RuntimeError: If WeasyPrint is not available
        """
        if not WEASYPRINT_AVAILABLE:
            raise RuntimeError(
                "PDF export requires WeasyPrint system dependencies. "
                "Install with: brew install pango (macOS) or "
                "apt-get install libpango-1.0-0 libpangocairo-1.0-0 (Ubuntu)"
            )

        if options is None:
            options = ExportOptions()

        full_html = self._wrap_html(html_content, options)

        # Render document to get page count without writing PDF
        document = HTML(string=full_html).render()
        return len(document.pages)

    def generate_docx(
        self,
        html_content: str,
        options: ExportOptions | None = None,
    ) -> bytes:
        """
        Generate DOCX from HTML content.

        Args:
            html_content: TipTap HTML content
            options: Export customization options

        Returns:
            DOCX file as bytes
        """
        if options is None:
            options = ExportOptions()

        doc = Document()

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(options.margin_top)
            section.bottom_margin = Inches(options.margin_bottom)
            section.left_margin = Inches(options.margin_left)
            section.right_margin = Inches(options.margin_right)

        # Set default style
        style = doc.styles["Normal"]
        style.font.name = options.font_family
        style.font.size = Pt(options.font_size)
        style.paragraph_format.line_spacing = options.line_spacing

        # Apply template-specific heading colors
        heading_colors = self._get_heading_colors(options.template)

        # Parse HTML and convert to DOCX
        self._parse_html_to_docx(doc, html_content, options, heading_colors)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _get_heading_colors(self, template: StyleTemplate) -> dict:
        """Get heading colors for the template."""
        if template == StyleTemplate.CLASSIC:
            return {
                "h1": RGBColor(0, 0, 0),
                "h2": RGBColor(51, 51, 51),
                "h3": RGBColor(68, 68, 68),
            }
        elif template == StyleTemplate.MODERN:
            return {
                "h1": RGBColor(26, 54, 93),
                "h2": RGBColor(37, 99, 235),
                "h3": RGBColor(30, 64, 175),
            }
        else:  # MINIMAL
            return {
                "h1": RGBColor(9, 9, 11),
                "h2": RGBColor(39, 39, 42),
                "h3": RGBColor(63, 63, 70),
            }

    def _parse_html_to_docx(
        self,
        doc: Document,
        html_content: str,
        options: ExportOptions,
        heading_colors: dict,
    ) -> None:
        """Parse HTML and add content to DOCX document."""
        # Clean and normalize HTML
        html_content = self._clean_html(html_content)

        # Split by major block elements
        # This is a simplified parser - for production, consider using BeautifulSoup
        blocks = self._split_html_blocks(html_content)

        for block in blocks:
            self._add_block_to_doc(doc, block, options, heading_colors)

    def _clean_html(self, html: str) -> str:
        """Clean and normalize HTML content."""
        # Remove extra whitespace
        html = re.sub(r'\s+', ' ', html)
        # Remove empty paragraphs
        html = re.sub(r'<p>\s*</p>', '', html)
        return html.strip()

    def _split_html_blocks(self, html: str) -> list[tuple[str, str]]:
        """
        Split HTML into blocks for processing.

        Returns list of (tag, content) tuples.
        """
        blocks = []

        # Pattern to match block-level elements
        block_pattern = re.compile(
            r'<(h[1-6]|p|ul|ol|blockquote)(?:\s[^>]*)?>(.+?)</\1>',
            re.IGNORECASE | re.DOTALL
        )

        pos = 0
        for match in block_pattern.finditer(html):
            # Check for text between matches
            if match.start() > pos:
                text_between = html[pos:match.start()].strip()
                if text_between:
                    # Treat as paragraph
                    blocks.append(("p", text_between))

            tag = match.group(1).lower()
            content = match.group(2)
            blocks.append((tag, content))
            pos = match.end()

        # Handle remaining text
        if pos < len(html):
            remaining = html[pos:].strip()
            if remaining:
                blocks.append(("p", remaining))

        return blocks

    def _add_block_to_doc(
        self,
        doc: Document,
        block: tuple[str, str],
        options: ExportOptions,
        heading_colors: dict,
    ) -> None:
        """Add a single block to the DOCX document."""
        tag, content = block

        if tag == "h1":
            para = doc.add_heading(level=1)
            self._add_formatted_text(para, content, options)
            para.runs[0].font.color.rgb = heading_colors["h1"]
            para.runs[0].font.size = Pt(options.font_size + 8)

        elif tag == "h2":
            para = doc.add_heading(level=2)
            self._add_formatted_text(para, content, options)
            para.runs[0].font.color.rgb = heading_colors["h2"]
            para.runs[0].font.size = Pt(options.font_size + 4)

        elif tag == "h3":
            para = doc.add_heading(level=3)
            self._add_formatted_text(para, content, options)
            para.runs[0].font.color.rgb = heading_colors["h3"]
            para.runs[0].font.size = Pt(options.font_size + 2)

        elif tag in ("ul", "ol"):
            self._add_list_to_doc(doc, content, tag == "ol", options)

        elif tag == "blockquote":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            self._add_formatted_text(para, content, options)
            para.runs[0].italic = True

        else:  # paragraph
            para = doc.add_paragraph()
            self._add_formatted_text(para, content, options)

    def _add_list_to_doc(
        self,
        doc: Document,
        content: str,
        is_ordered: bool,
        options: ExportOptions,
    ) -> None:
        """Add a list to the DOCX document."""
        # Extract list items
        li_pattern = re.compile(r'<li>(.+?)</li>', re.IGNORECASE | re.DOTALL)
        items = li_pattern.findall(content)

        for item in items:
            style = "List Number" if is_ordered else "List Bullet"
            para = doc.add_paragraph(style=style)
            self._add_formatted_text(para, item, options)

    def _add_formatted_text(
        self,
        paragraph,
        html_text: str,
        options: ExportOptions,
    ) -> None:
        """Add formatted text to a paragraph, handling inline HTML tags."""
        # Pattern to find inline formatting
        parts = self._split_inline_formatting(html_text)

        for text, styles in parts:
            # Strip any remaining HTML tags for the text
            clean_text = re.sub(r'<[^>]+>', '', text)
            if not clean_text:
                continue

            run = paragraph.add_run(clean_text)
            run.font.name = options.font_family
            run.font.size = Pt(options.font_size)

            if "bold" in styles:
                run.bold = True
            if "italic" in styles:
                run.italic = True
            if "underline" in styles:
                run.underline = True
            if "strike" in styles:
                run.font.strike = True

    def _split_inline_formatting(self, html: str) -> list[tuple[str, set]]:
        """
        Split text into segments with their formatting.

        Returns list of (text, set_of_styles) tuples.
        """
        result = []

        # Stack-based parser for nested tags
        def parse_recursive(text: str, current_styles: set) -> None:
            # Find the first tag
            tag_match = re.search(r'<(strong|b|em|i|u|s|strike)>(.+?)</\1>', text, re.IGNORECASE | re.DOTALL)

            if not tag_match:
                # No more tags, add remaining text
                if text.strip():
                    result.append((text, current_styles.copy()))
                return

            # Add text before the tag
            before = text[:tag_match.start()]
            if before.strip():
                result.append((before, current_styles.copy()))

            # Determine the style
            tag = tag_match.group(1).lower()
            new_styles = current_styles.copy()
            if tag in ("strong", "b"):
                new_styles.add("bold")
            elif tag in ("em", "i"):
                new_styles.add("italic")
            elif tag == "u":
                new_styles.add("underline")
            elif tag in ("s", "strike"):
                new_styles.add("strike")

            # Recursively parse the content inside the tag
            parse_recursive(tag_match.group(2), new_styles)

            # Continue with text after the tag
            after = text[tag_match.end():]
            if after:
                parse_recursive(after, current_styles)

        parse_recursive(html, set())
        return result

    # Async wrapper methods for API compatibility
    async def export_pdf(
        self,
        html_content: str,
        template: str = "classic",
        font_family: str = "Arial",
        font_size: int = 11,
        margin_top: float = 0.75,
        margin_bottom: float = 0.75,
        margin_left: float = 0.75,
        margin_right: float = 0.75,
    ) -> bytes:
        """Export HTML content as PDF bytes."""
        try:
            style_template = StyleTemplate(template)
        except ValueError:
            style_template = StyleTemplate.CLASSIC

        options = ExportOptions(
            template=style_template,
            font_family=font_family,
            font_size=font_size,
            margin_top=margin_top,
            margin_bottom=margin_bottom,
            margin_left=margin_left,
            margin_right=margin_right,
        )
        return self.generate_pdf(html_content, options)

    async def export_docx(
        self,
        html_content: str,
        template: str = "classic",
        font_family: str = "Arial",
        font_size: int = 11,
        margin_top: float = 0.75,
        margin_bottom: float = 0.75,
        margin_left: float = 0.75,
        margin_right: float = 0.75,
    ) -> bytes:
        """Export HTML content as DOCX bytes."""
        try:
            style_template = StyleTemplate(template)
        except ValueError:
            style_template = StyleTemplate.CLASSIC

        options = ExportOptions(
            template=style_template,
            font_family=font_family,
            font_size=font_size,
            margin_top=margin_top,
            margin_bottom=margin_bottom,
            margin_left=margin_left,
            margin_right=margin_right,
        )
        return self.generate_docx(html_content, options)


# Singleton instance
_html_export_service: HTMLToDocumentService | None = None


def get_html_export_service() -> HTMLToDocumentService:
    """Get the HTML export service singleton."""
    global _html_export_service
    if _html_export_service is None:
        _html_export_service = HTMLToDocumentService()
    return _html_export_service
